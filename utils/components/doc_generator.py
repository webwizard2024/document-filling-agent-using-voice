# components/doc_generator.py
import io
from docx import Document

class DocumentGenerator:
    """Generates a downloadable DOCX file from text."""

    def create_docx(self, content: str, title: str = "Generated Document") -> io.BytesIO:
        """Creates a DOCX file in memory."""
        doc = Document()
        
        # Add a title to the document
        doc.add_heading(title, level=1)
        
        # Add the content as a paragraph
        doc.add_paragraph(content)
        
        # Save the document to an in-memory bytes buffer
        doc_buffer = io.BytesIO()
        doc.save(doc_buffer)
        doc_buffer.seek(0)
        
        return doc_buffer